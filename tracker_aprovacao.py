#!/usr/bin/env python3
"""
Bot de Rastreamento de Concursos ETEC/FATEC
Crawler autônomo que descobre processos seletivos no portal do CPS (URH),
baixa documentos publicados e verifica se o nome do candidato aparece.

Também monitora o Diário Oficial do Estado de SP (DOE SP) buscando
citações do nome do candidato em publicações oficiais.

Extrai metadados (edital, unidade, cidade, disciplina) e identifica
a fase do processo (Abertura → Classificação → Convocação…).

Notificações via WhatsApp (CallMeBot).
"""

import io
import json
import os
import re
import time
from datetime import datetime, timedelta
from pathlib import Path
from urllib.parse import urljoin

import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

# ──────────────────────────────────────────────
# CONFIGURAÇÃO
# ──────────────────────────────────────────────

MEU_NOME = os.getenv("MEU_NOME", "Renan Bezerra dos Santos")

# CallMeBot – WhatsApp
CALLMEBOT_PHONE = os.getenv("CALLMEBOT_PHONE", "")
CALLMEBOT_APIKEY = os.getenv("CALLMEBOT_APIKEY", "")

# Caminho do histórico de documentos já processados
HISTORY_FILE = Path(__file__).parent / "history_pdfs.json"

# Headers para simular navegador comum
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    )
}

REQUEST_TIMEOUT = 30  # segundos

# Base do portal CPS
CPS_BASE = "https://urhsistemas.cps.sp.gov.br"

# Páginas de listagem de processos (Inscrições Abertas + Em Andamento)
LISTING_PAGES: list[dict] = [
    # ── ETEC ──
    {"url": f"{CPS_BASE}/dgsdad/SelecaoPublica/ETEC/PSS/Abertos.aspx",
     "label": "ETEC – Processo Seletivo Docente – Inscrições Abertas"},
    {"url": f"{CPS_BASE}/dgsdad/SelecaoPublica/ETEC/PSS/Andamento.aspx",
     "label": "ETEC – Processo Seletivo Docente – Em Andamento"},
    {"url": f"{CPS_BASE}/dgsdad/selecaopublica/ETEC/CPD/Abertos.aspx",
     "label": "ETEC – Concurso Público Docente – Inscrições Abertas"},
    {"url": f"{CPS_BASE}/dgsdad/selecaopublica/ETEC/CPD/emAndamento.aspx",
     "label": "ETEC – Concurso Público Docente – Em Andamento"},
    {"url": f"{CPS_BASE}/dgsdad/SelecaoPublica/ETEC/Auxiliar/EmAndamento.aspx",
     "label": "ETEC – Auxiliar de Docente – Em Andamento"},
    # ── FATEC ──
    {"url": f"{CPS_BASE}/dgsdad/SelecaoPublica/FATEC/PSS/inscricoesabertas.aspx",
     "label": "FATEC – Processo Seletivo Docente – Inscrições Abertas"},
    {"url": f"{CPS_BASE}/dgsdad/SelecaoPublica/FATEC/ProcessoSeletivo/EmAndamento.aspx",
     "label": "FATEC – Processo Seletivo Docente – Em Andamento"},
    {"url": f"{CPS_BASE}/dgsdad/SelecaoPublica/FATEC/CPD/Abertos.aspx",
     "label": "FATEC – Concurso Público Docente – Inscrições Abertas"},
    {"url": f"{CPS_BASE}/dgsdad/SelecaoPublica/FATEC/CPD/emAndamento.aspx",
     "label": "FATEC – Concurso Público Docente – Em Andamento"},
    # ── PSSAD (Auxiliar de Docente – compartilhado ETEC/FATEC) ──
    {"url": f"{CPS_BASE}/dgsdad/SelecaoPublica/PSSAD/Abertos.aspx",
     "label": "PSSAD – Auxiliar de Docente – Inscrições Abertas"},
    {"url": f"{CPS_BASE}/dgsdad/selecaopublica/PSSAD/emAndamento.aspx",
     "label": "PSSAD – Auxiliar de Docente – Em Andamento"},
]

# Limite de processos por página de listagem (evita sobrecarga)
MAX_PROCESSES_PER_PAGE = 50

# ──────────────────────────────────────────────
# MAPA DE FASES – classifica o documento pelo nome/link
# Ordem reflete a progressão real do processo seletivo.
# ──────────────────────────────────────────────

PHASE_MAP: list[tuple[str, str]] = [
    # (padrão regex case-insensitive, rótulo amigável)
    (r"edital\s*de\s*abertura|EDITALDE?ABERTURA", "📋 Edital de Abertura"),
    (r"redu[cç][aã]o.*isen[cç][aã]o|REDUCAO.ISENCAO", "💰 Resultado Redução/Isenção de Taxa"),
    (r"reabertura|REABERTURA", "🔄 Reabertura de Inscrições"),
    (r"banca\s*examinadora|BANCAEXAMINADORA", "👥 Portaria da Banca Examinadora"),
    (r"altera[cç][aã]o.*cronograma|ALTERACAOCRONOGRAMA", "📅 Alteração de Cronograma"),
    (r"altera[cç][aã]o.*comiss[aã]o|ALTERACAOCOMISSAO", "🔀 Alteração da Comissão"),
    (r"deferimento|indeferimento|DEFERIMENTO", "✅ Deferimento/Indeferimento de Inscrições"),
    (r"resultado.*escrita.*conv|RESULTADOESCRITACONV|resultado.*pve", "📝 Resultado Prova Escrita e Convocação Didática"),
    (r"resultado.*memorial|resultado.*prova|RESULTADO", "📝 Resultado de Prova/Avaliação"),
    (r"classifica[cç][aã]o\s*final|CLASSIFICAOFINAL|CLASSIFICACAOFINAL", "🏆 Classificação Final"),
    (r"homologa[cç][aã]o|HOMOLOGA", "✔️ Homologação"),
    (r"convoca[cç][aã]o|CONVOCAO|CONVOCACAO", "📞 Convocação"),
    (r"prorroga[cç][aã]o|PRORROGA", "⏳ Prorrogação de Validade"),
]


def classify_phase(doc_name: str, doc_url: str) -> str:
    """Identifica a fase do processo a partir do nome do documento ou URL."""
    combined = f"{doc_name} {doc_url}"
    for pattern, phase_label in PHASE_MAP:
        if re.search(pattern, combined, re.IGNORECASE):
            return phase_label
    return "📄 Documento"


# ──────────────────────────────────────────────
# DIÁRIO OFICIAL DO ESTADO DE SP (DOE SP)
# API pública: do-api-web-search.doe.sp.gov.br
# ──────────────────────────────────────────────

DOE_API_BASE = "https://do-api-web-search.doe.sp.gov.br"
DOE_SITE_BASE = "https://www.doe.sp.gov.br"

# ID do caderno "Executivo" no DOE SP
DOE_JOURNAL_EXECUTIVO = "ca96256b-6ca1-407f-866e-567ef9430123"

# Quantos dias para trás buscar no DOE (janela de busca)
DOE_SEARCH_DAYS = 30

# Máximo de resultados por página na API do DOE
DOE_PAGE_SIZE = 20


def search_doe_sp(name: str, history: dict) -> tuple[dict, int]:
    """
    Busca o nome do candidato no Diário Oficial do Estado de SP
    via API pública. Retorna (history atualizado, qtd novos).
    """
    today = datetime.now()
    from_date = (today - timedelta(days=DOE_SEARCH_DAYS)).strftime("%Y-%m-%d")
    to_date = today.strftime("%Y-%m-%d")

    new_count = 0
    page = 1

    while True:
        params = {
            "Terms": name,
            "FromDate": from_date,
            "ToDate": to_date,
            "JournalId": DOE_JOURNAL_EXECUTIVO,
            "PageNumber": page,
            "PageSize": DOE_PAGE_SIZE,
            "SortField": "Date",
        }

        try:
            resp = requests.get(
                f"{DOE_API_BASE}/v2/advanced-search/publications",
                params=params,
                headers=HEADERS,
                timeout=REQUEST_TIMEOUT,
            )
            resp.raise_for_status()
            data = resp.json()
        except (requests.RequestException, ValueError) as e:
            print(f"  [ERRO] Falha na busca DOE SP (página {page}): {e}")
            break

        items = data.get("items", [])
        if not items:
            break

        for item in items:
            pub_id = item.get("id", "")
            doe_key = f"doe:{pub_id}"

            if doe_key in history:
                continue

            title = item.get("title", "Sem título")
            slug = item.get("slug", "")
            hierarchy = item.get("hierarchy", "")
            excerpt = item.get("excerpt", "")
            pub_date = item.get("date", "")[:10]
            pub_url = f"{DOE_SITE_BASE}/{slug}" if slug else ""
            matches = item.get("totalTermsFound", 0)

            print(f"    [DOE NOVO] {title}")
            print(f"      Hierarquia: {hierarchy}")
            print(f"      Menções: {matches}")

            history[doe_key] = {
                "source": "DOE-SP",
                "title": title,
                "date": pub_date,
                "hierarchy": hierarchy,
                "url": pub_url,
                "matches": matches,
                "found_name": True,
            }
            new_count += 1

            # Montar mensagem WhatsApp
            msg = (
                "📰 *SEU NOME NO DIÁRIO OFICIAL!* 📰\n\n"
                f"📌 *Publicação:* {title}\n"
                f"📅 *Data:* {pub_date}\n"
                f"🏛️ *Seção:* {hierarchy}\n"
                f"🔎 *Menções encontradas:* {matches}\n"
            )
            if excerpt:
                # Limitar excerpt para não estourar mensagem
                short_excerpt = excerpt[:300]
                if len(excerpt) > 300:
                    short_excerpt += "…"
                msg += f"📝 *Trecho:* _{short_excerpt}_\n"
            if pub_url:
                msg += f"🔗 *Link:* {pub_url}"

            send_whatsapp(msg)

        # Próxima página
        if not data.get("hasNextPage", False):
            break
        page += 1

    return history, new_count


# ──────────────────────────────────────────────
# HISTÓRICO
# ──────────────────────────────────────────────

def load_history() -> dict:
    """Carrega o JSON com os documentos já processados."""
    if HISTORY_FILE.exists():
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_history(history: dict) -> None:
    """Salva o JSON atualizado."""
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


# ──────────────────────────────────────────────
# NOTIFICAÇÃO – WHATSAPP (CallMeBot)
# ──────────────────────────────────────────────

def send_whatsapp(message: str) -> None:
    """Envia mensagem via CallMeBot WhatsApp API."""
    if not CALLMEBOT_PHONE or not CALLMEBOT_APIKEY:
        print("[AVISO] CallMeBot não configurado. Mensagem apenas no log:")
        print(message)
        return

    url = "https://api.callmebot.com/whatsapp.php"
    params = {
        "phone": CALLMEBOT_PHONE,
        "text": message,
        "apikey": CALLMEBOT_APIKEY,
    }
    try:
        resp = requests.get(url, params=params, timeout=REQUEST_TIMEOUT)
        if resp.status_code == 200:
            print("[OK] Mensagem WhatsApp enviada.")
        else:
            print(f"[ERRO] CallMeBot retornou status {resp.status_code}: {resp.text}")
    except requests.RequestException as e:
        print(f"[ERRO] Falha ao enviar WhatsApp: {e}")

    # Respeitar rate-limit do CallMeBot (mín. 2 s entre mensagens)
    time.sleep(3)


# ──────────────────────────────────────────────
# CRAWLER – DESCOBERTA DE PROCESSOS
# ──────────────────────────────────────────────

def _get_soup(url: str) -> BeautifulSoup | None:
    """Faz GET e retorna BeautifulSoup ou None em caso de erro."""
    try:
        resp = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT)
        resp.raise_for_status()
        return BeautifulSoup(resp.text, "html.parser")
    except requests.RequestException as e:
        print(f"  [ERRO] Não foi possível acessar {url}: {e}")
        return None


def discover_detail_links(listing_url: str) -> list[str]:
    """
    Acessa uma página de listagem (GridView) e extrai os links
    das páginas de detalhes dos processos seletivos.
    Retorna até MAX_PROCESSES_PER_PAGE URLs únicas.
    """
    soup = _get_soup(listing_url)
    if soup is None:
        return []

    detail_links: list[str] = []
    seen: set[str] = set()

    for a_tag in soup.find_all("a", href=True):
        href = a_tag["href"]
        # Páginas de detalhes contêm o parâmetro oljioahohafnav87412
        if "oljioahohafnav87412" not in href:
            continue
        # Ignorar links javascript:__doPostBack (são os cabeçalhos de ordenação)
        if href.startswith("javascript:"):
            continue
        full_url = urljoin(listing_url, href)
        if full_url not in seen:
            seen.add(full_url)
            detail_links.append(full_url)

    return detail_links[:MAX_PROCESSES_PER_PAGE]


# ──────────────────────────────────────────────
# EXTRAÇÃO DE METADADOS DA PÁGINA DE DETALHES
# ──────────────────────────────────────────────

def extract_metadata(soup: BeautifulSoup) -> dict:
    """
    Extrai metadados do processo a partir do texto da página de detalhes.
    Retorna dict com: edital, unidade, cidade, disciplina/curso.
    """
    text = soup.get_text(" ", strip=True)

    meta: dict = {
        "edital": "",
        "unidade": "",
        "cidade": "",
        "disciplina": "",
    }

    # Nº do Edital – ex: "EDITAL DE ABERTURA Nº  229/11/2026"
    m = re.search(
        r"EDITAL\s+DE\s+ABERTURA\s+N[ºo°]\s*([\d/]+)",
        text, re.IGNORECASE,
    )
    if m:
        meta["edital"] = m.group(1).strip()

    # Unidade de Ensino e Cidade
    # Padrão: "CÓD. DA UNIDADE:  229 - UNIDADE DE ENSINO:  Escola ... - CIDADE: São Paulo"
    m = re.search(
        r"UNIDADE\s+DE\s+ENSINO:\s*(.+?)\s*-\s*CIDADE:\s*(.+?)(?:\n|CURSO|DISCIPLINA|COMPONENTE|REQUISITO|Os pedidos|Per[ií]odo)",
        text, re.IGNORECASE,
    )
    if m:
        meta["unidade"] = m.group(1).strip()
        meta["cidade"] = m.group(2).strip()

    # Disciplina ou Componente Curricular
    m = re.search(
        r"(?:DISCIPLINA|COMPONENTE\s+CURRICULAR):\s*(?:\d+\s*-\s*)?(.+?)(?:\n|REQUISITO|Os pedidos|Per[ií]odo)",
        text, re.IGNORECASE,
    )
    if m:
        meta["disciplina"] = m.group(1).strip()

    # Se não achou disciplina, tenta CURSO
    if not meta["disciplina"]:
        m = re.search(
            r"CURSO:\s*(.+?)(?:\n|DISCIPLINA|COMPONENTE|REQUISITO|Os pedidos|Per[ií]odo)",
            text, re.IGNORECASE,
        )
        if m:
            meta["disciplina"] = m.group(1).strip()

    return meta


# ──────────────────────────────────────────────
# SCRAPING – DOCUMENTOS NA PÁGINA DE DETALHES
# ──────────────────────────────────────────────

def fetch_detail_page(detail_url: str) -> tuple[dict, list[dict]]:
    """
    Acessa a página de detalhes de um processo e retorna:
    1. Metadados (edital, unidade, cidade, disciplina)
    2. Lista de documentos encontrados (PDF e DOCX)
    """
    soup = _get_soup(detail_url)
    if soup is None:
        return {}, []

    meta = extract_metadata(soup)

    docs: list[dict] = []
    doc_pattern = re.compile(r"\.(pdf|docx?)(\?.*)?$", re.IGNORECASE)

    for a_tag in soup.find_all("a", href=True):
        href = a_tag["href"]
        match = doc_pattern.search(href)
        if not match:
            continue
        full_url = urljoin(detail_url, href)
        name = a_tag.get_text(strip=True) or href.split("/")[-1].split("?")[0]
        ext = match.group(1).lower()
        if ext == "doc":
            ext = "docx"
        phase = classify_phase(name, full_url)
        docs.append({"name": name, "url": full_url, "ext": ext, "phase": phase})

    return meta, docs


# ──────────────────────────────────────────────
# ANÁLISE DE DOCUMENTOS (PDF e DOCX)
# ──────────────────────────────────────────────

def _download(url: str) -> bytes | None:
    """Baixa um arquivo na memória e retorna os bytes."""
    try:
        resp = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT * 2)
        resp.raise_for_status()
        return resp.content
    except requests.RequestException as e:
        print(f"  [ERRO] Falha ao baixar {url}: {e}")
        return None


def check_name_in_pdf(content: bytes, name: str) -> bool:
    """Verifica se o nome aparece em um PDF (case insensitive)."""
    name_lower = name.lower()
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if name_lower in text.lower():
                    return True
    except Exception as e:
        print(f"  [ERRO] Falha ao ler PDF: {e}")
    return False


def check_name_in_docx(content: bytes, name: str) -> bool:
    """Verifica se o nome aparece em um DOCX (case insensitive)."""
    name_lower = name.lower()
    try:
        doc = DocxDocument(io.BytesIO(content))
        for para in doc.paragraphs:
            if name_lower in para.text.lower():
                return True
        # Verificar também tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if name_lower in cell.text.lower():
                        return True
    except Exception as e:
        print(f"  [ERRO] Falha ao ler DOCX: {e}")
    return False


def check_name_in_document(url: str, ext: str, name: str) -> bool:
    """Baixa o documento e verifica se o nome aparece."""
    content = _download(url)
    if content is None:
        return False
    if ext == "pdf":
        return check_name_in_pdf(content, name)
    elif ext in ("docx", "doc"):
        return check_name_in_docx(content, name)
    return False


# ──────────────────────────────────────────────
# FORMATAÇÃO DE MENSAGENS
# ──────────────────────────────────────────────

def _format_meta_block(meta: dict, label: str) -> str:
    """Monta o bloco de informações do processo para a mensagem."""
    lines: list[str] = []
    if meta.get("edital"):
        lines.append(f"📌 *Edital:* {meta['edital']}")
    if meta.get("unidade"):
        lines.append(f"🏫 *Unidade:* {meta['unidade']}")
    if meta.get("cidade"):
        lines.append(f"📍 *Cidade:* {meta['cidade']}")
    if meta.get("disciplina"):
        lines.append(f"📚 *Disciplina:* {meta['disciplina']}")
    lines.append(f"🗂️ *Tipo:* {label}")
    return "\n".join(lines)


def build_message_found(doc_name: str, doc_url: str, phase: str,
                        meta: dict, label: str, detail_url: str) -> str:
    """Mensagem quando o nome É encontrado no documento."""
    return (
        "🚨🚨🚨 *SEU NOME FOI ENCONTRADO!* 🚨🚨🚨\n\n"
        f"{_format_meta_block(meta, label)}\n\n"
        f"📄 *Documento:* {doc_name}\n"
        f"🔖 *Fase:* {phase}\n"
        f"🔗 *Arquivo:* {doc_url}\n"
        f"📋 *Página:* {detail_url}"
    )


def build_message_not_found(doc_name: str, doc_url: str, phase: str,
                            meta: dict, label: str, detail_url: str) -> str:
    """Mensagem quando o nome NÃO é encontrado (nova publicação)."""
    return (
        "⚠️ *Nova publicação detectada*\n\n"
        f"{_format_meta_block(meta, label)}\n\n"
        f"📄 *Documento:* {doc_name}\n"
        f"🔖 *Fase:* {phase}\n"
        "Seu nome *não* foi encontrado na busca automática.\n"
        f"🔗 *Arquivo:* {doc_url}\n"
        f"📋 *Página:* {detail_url}"
    )


# ──────────────────────────────────────────────
# LÓGICA PRINCIPAL
# ──────────────────────────────────────────────

def process_detail_page(detail_url: str, label: str, history: dict) -> dict:
    """
    Processa uma página de detalhes de um processo seletivo.
    Retorna o history atualizado.
    """
    meta, docs = fetch_detail_page(detail_url)
    if not docs:
        return history

    edital_info = meta.get("edital", "?")
    unidade_info = meta.get("unidade", "?")

    for doc_info in docs:
        doc_url = doc_info["url"]
        doc_name = doc_info["name"]
        doc_ext = doc_info["ext"]
        doc_phase = doc_info["phase"]

        if doc_url in history:
            continue

        print(f"    [NOVO] {doc_phase} | {doc_name} (.{doc_ext})")
        found = check_name_in_document(doc_url, doc_ext, MEU_NOME)

        # Registrar no histórico
        history[doc_url] = {
            "name": doc_name,
            "phase": doc_phase,
            "detail_page": detail_url,
            "listing": label,
            "edital": meta.get("edital", ""),
            "unidade": meta.get("unidade", ""),
            "cidade": meta.get("cidade", ""),
            "disciplina": meta.get("disciplina", ""),
            "found_name": found,
        }

        if found:
            msg = build_message_found(
                doc_name, doc_url, doc_phase, meta, label, detail_url)
            print(f"    >>> NOME ENCONTRADO! Edital {edital_info} – {unidade_info} <<<")
        else:
            msg = build_message_not_found(
                doc_name, doc_url, doc_phase, meta, label, detail_url)
            print(f"    Nome não encontrado. Edital {edital_info} – {unidade_info}")

        send_whatsapp(msg)

    return history


def main() -> None:
    print("Bot de Rastreamento de Concursos ETEC/FATEC (Crawler Autônomo)")
    print(f"Nome monitorado: {MEU_NOME}")
    print(f"Páginas de listagem: {len(LISTING_PAGES)}")

    history = load_history()
    total_new = 0

    for listing in LISTING_PAGES:
        listing_url = listing["url"]
        label = listing["label"]

        print(f"\n{'='*60}")
        print(f"[LISTAGEM] {label}")
        print(f"  {listing_url}")
        print(f"{'='*60}")

        detail_links = discover_detail_links(listing_url)
        if not detail_links:
            print("  Nenhum processo encontrado nesta página.")
            continue

        print(f"  {len(detail_links)} processo(s) encontrado(s).")

        for i, detail_url in enumerate(detail_links, 1):
            print(f"  [{i}/{len(detail_links)}] {detail_url}")
            old_count = len(history)
            history = process_detail_page(detail_url, label, history)
            total_new += len(history) - old_count

    # ── FASE 2: Diário Oficial do Estado de SP ──
    print(f"\n{'='*60}")
    print(f"[DOE SP] Buscando nome no Diário Oficial do Estado de SP")
    print(f"  Período: últimos {DOE_SEARCH_DAYS} dias")
    print(f"{'='*60}")

    history, doe_new = search_doe_sp(MEU_NOME, history)
    total_new += doe_new

    if doe_new == 0:
        print("  Nenhuma publicação nova encontrada no DOE SP.")
    else:
        print(f"  {doe_new} publicação(ões) nova(s) no DOE SP.")

    save_history(history)
    print(f"\n{'='*60}")
    print(f"Execução finalizada.")
    print(f"  Documentos novos processados: {total_new}")
    print(f"  Total no histórico: {len(history)}")
    print(f"  Histórico salvo em {HISTORY_FILE}")


if __name__ == "__main__":
    main()
